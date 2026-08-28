"""
doc_generator.py
Generates BRD.docx and TDD.docx from canonical JSON using python-docx.
Template structure is hardcoded — LLM cannot affect formatting.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import io


# ── Helpers ──
def v(field):
    if field is None:
        return "TBD"
    if isinstance(field, str):
        return field
    if isinstance(field, dict) and "value" in field:
        val = str(field["value"])
        if field.get("tag") == "ASSUMPTION" and not val.endswith("[ASSUMPTION]"):
            return val + " [ASSUMPTION]"
        return val
    return "TBD"


def v_raw(field):
    """Get value without ASSUMPTION suffix."""
    if field is None:
        return "TBD"
    if isinstance(field, str):
        return field
    if isinstance(field, dict) and "value" in field:
        return str(field["value"])
    return "TBD"


HEADER_COLOR = "1F3864"
HEADER_TEXT = "FFFFFF"
ALT_ROW = "F2F2F2"
BORDER_COLOR = "B4C6E7"
FONT_NAME = "Calibri"
FONT_SIZE = Pt(10)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_header_row(table, texts):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = FONT_SIZE
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor.from_string(HEADER_TEXT)
        set_cell_shading(cell, HEADER_COLOR)


def add_data_row(table, values, shaded=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = FONT_SIZE
        run.font.name = FONT_NAME
        if shaded:
            set_cell_shading(cell, ALT_ROW)
    return row


def add_label_value_row(table, label, value, shaded=False):
    row = table.add_row()
    # Label cell
    cell0 = row.cells[0]
    cell0.text = ""
    run0 = cell0.paragraphs[0].add_run(label)
    run0.bold = True
    run0.font.size = FONT_SIZE
    run0.font.name = FONT_NAME
    # Value cell
    cell1 = row.cells[1]
    cell1.text = ""
    run1 = cell1.paragraphs[0].add_run(str(value))
    run1.font.size = FONT_SIZE
    run1.font.name = FONT_NAME
    if shaded:
        set_cell_shading(cell0, ALT_ROW)
        set_cell_shading(cell1, ALT_ROW)
    return row


def two_col_table(doc, rows_data):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    # Set header for first row
    cell0 = table.rows[0].cells[0]
    cell1 = table.rows[0].cells[1]
    cell0.text = ""
    cell1.text = ""
    r0 = cell0.paragraphs[0].add_run(rows_data[0][0])
    r0.bold = True
    r0.font.size = FONT_SIZE
    r0.font.name = FONT_NAME
    r1 = cell1.paragraphs[0].add_run(str(rows_data[0][1]))
    r1.font.size = FONT_SIZE
    r1.font.name = FONT_NAME

    for i, (label, value) in enumerate(rows_data[1:], 1):
        add_label_value_row(table, label, value, shaded=(i % 2 == 0))

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(HEADER_COLOR)
    return h


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = FONT_SIZE
    run.font.name = FONT_NAME
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(str(item))
        run.font.size = FONT_SIZE
        run.font.name = FONT_NAME


# ════════════════════════════════════════
#  BRD GENERATOR
# ════════════════════════════════════════
def generate_brd(c):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    project_name = v(c.get("project", {}).get("name"))
    opco = v(c.get("project", {}).get("opco"))
    date_str = v(c.get("project", {}).get("date", {"value": datetime.now().strftime("%b %d, %Y")}))
    wt_code = v_raw(c.get("integration", {}).get("workTypeCode"))

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Business Requirements Document")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(HEADER_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Integration Projects")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(HEADER_COLOR)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"{project_name} | {opco} | {date_str}")
    run.bold = True
    run.font.size = FONT_SIZE

    # Version table
    vtable = doc.add_table(rows=2, cols=4)
    vtable.style = "Table Grid"
    for i, text in enumerate(["Version", v(c["project"].get("version")), "Author", v(c["project"].get("author"))]):
        cell = vtable.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.size = FONT_SIZE
        run.font.name = FONT_NAME
        if i % 2 == 0:
            run.bold = True
    for i, text in enumerate(["Date", date_str, "Reviewer", v(c["project"].get("reviewer"))]):
        cell = vtable.rows[1].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.size = FONT_SIZE
        run.font.name = FONT_NAME
        if i % 2 == 0:
            run.bold = True
        set_cell_shading(cell, ALT_ROW)

    doc.add_page_break()

    # ── Project Overview ──
    add_heading(doc, "Project Overview", level=1)

    def wt_check(code):
        return "☑" if wt_code == code else "☐"

    wt_text = (
        f"{wt_check('T1')} T1 – File to Database\n"
        f"{wt_check('T2')} T2 – API to Database\n"
        f"{wt_check('T3')} T3 – Database to Database\n"
        f"{wt_check('T4')} T4 – Database to File\n"
        f"{wt_check('T5')} T5 – Database to API\n"
        f"Integration Type: {v(c.get('integration', {}).get('direction'))}"
    )

    classification = v_raw(c.get("project", {}).get("classification"))
    cls_text = (
        f"{'☑' if classification == 'New Project' else '☐'} New Project  "
        f"{'☑' if classification == 'Phase of Existing Project' else '☐'} Phase of Existing Project  "
        f"{'☑' if classification == 'Enhancement' else '☐'} Enhancement"
    )

    entities = ", ".join(v(e) for e in c.get("entities", [{"value": "TBD"}]))

    two_col_table(doc, [
        ["Project Name", project_name],
        ["OpCo Name", opco],
        ["Business Function", v(c.get("project", {}).get("businessFunction"))],
        ["Business Contacts", v(c.get("project", {}).get("businessContacts"))],
        ["Brief Description", v(c.get("requirement", {}).get("businessObjective"))],
        ["Project Classification", cls_text],
        ["Type of Work", wt_text],
        ["Source System(s)", v(c.get("source", {}).get("system"))],
        ["Destination System(s)", v(c.get("target", {}).get("system"))],
    ])

    doc.add_paragraph()

    # ── Customer Requirement Statement ──
    add_heading(doc, "Customer Requirement Statement", level=1)

    scope = "\n".join("• " + v(a) for a in c.get("acceptanceCriteria", []))
    out_of_scope = "\n".join("• " + v(a) for a in c.get("outOfScope", []))
    acceptance = "\n".join("• " + v(a) for a in c.get("acceptanceCriteria", []))

    two_col_table(doc, [
        ["Requirement Statement", v(c.get("requirement", {}).get("statement"))],
        ["Visual Flow Diagrams", "TBD — to be created during detailed design"],
        ["Data Entities", entities],
        ["Mapping Document", v(c.get("mappings", {}).get("status"))],
        ["Integration Platform", v(c.get("integration", {}).get("middleware"))],
        ["Scope Information", scope or "TBD"],
        ["Out of Scope", out_of_scope or "TBD"],
        ["Acceptance Criteria", acceptance or "TBD"],
        ["Reference Integrations", "N/A — new integration"],
        ["Detailed Estimate", "TBD — to be estimated after detailed design"],
    ])

    doc.add_paragraph()

    # ── Unresolved Items ──
    add_heading(doc, "Unresolved Items / TBDs", level=1)
    add_para(doc, "The following items require clarification or confirmation before or during detailed design:")
    add_bullets(doc, c.get("unresolvedItems", ["No unresolved items identified"]))

    doc.add_paragraph()

    # ── Change Log ──
    add_heading(doc, "Change Log", level=1)
    add_para(doc, "Track changes or additions to previously stated requirements.")

    cl_table = doc.add_table(rows=1, cols=4)
    cl_table.style = "Table Grid"
    make_header_row(cl_table, ["Date", "Changeset ID", "Requestor", "Description of Changes"])
    add_data_row(cl_table, [date_str, "CS-001", "AI-DocGen", "Initial Version — BRD auto-generated from requirement."])

    doc.add_paragraph()
    add_para(doc, f"Confidential – {opco}")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ════════════════════════════════════════
#  TDD GENERATOR
# ════════════════════════════════════════
def generate_tdd(c):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    project_name = v(c.get("project", {}).get("name"))
    opco = v(c.get("project", {}).get("opco"))
    date_str = v(c.get("project", {}).get("date", {"value": datetime.now().strftime("%b %d, %Y")}))
    src = v(c.get("source", {}).get("system"))
    tgt = v(c.get("target", {}).get("system"))
    middleware = v(c.get("integration", {}).get("middleware"))
    pattern = v(c.get("integration", {}).get("pattern"))
    entities = ", ".join(v(e) for e in c.get("entities", [{"value": "TBD"}]))
    entities_lower = entities.lower()
    biz_func = v(c.get("project", {}).get("businessFunction"))
    protocol = v(c.get("source", {}).get("protocol"))

    # ── Title Page ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Enterprise Boomi Technical Design Document (TDD)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(HEADER_COLOR)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(project_name)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(HEADER_COLOR)

    # Document info
    add_heading(doc, "Document Information", level=2)
    two_col_table(doc, [
        ["Document Name", "Technical Design Document"],
        ["Project Name", project_name],
        ["Integration Name", project_name],
        ["OPU / OPCO", opco],
        ["Version", v(c["project"].get("version"))],
        ["Author", v(c["project"].get("author"))],
        ["Reviewer", v(c["project"].get("reviewer"))],
        ["Approver", "TBD"],
        ["Creation Date", date_str],
        ["Last Modified Date", date_str],
    ])

    doc.add_paragraph()
    add_heading(doc, "Revision History", level=2)
    rh = doc.add_table(rows=1, cols=4)
    rh.style = "Table Grid"
    make_header_row(rh, ["Version", "Date", "Author", "Description"])
    add_data_row(rh, ["1.0", date_str, "AI-DocGen", "Initial TDD auto-generated from requirement."])

    doc.add_page_break()

    # ── Section 1: Introduction ──
    add_heading(doc, "1. Introduction", level=1)
    add_heading(doc, "1.1 Purpose", level=2)
    add_para(doc, f"This document defines the technical design for the integration that processes {entities_lower} data from {src} and loads it into {tgt}. It covers architecture, process flow, mappings, validations, error handling, deployment, monitoring, and support procedures.")

    add_heading(doc, "1.2 Business Background", level=2)
    add_para(doc, v(c.get("requirement", {}).get("businessObjective")))

    add_heading(doc, "1.3 Objectives", level=2)
    add_bullets(doc, [
        f"Automate the ingestion of {entities_lower} data from {src} into {tgt}.",
        "Ensure accurate data transformation and validation before loading.",
        "Minimize manual intervention in the data transfer process.",
        f"Improve operational efficiency for the {biz_func.lower()} function.",
    ])

    add_heading(doc, "1.4 Success Criteria", level=2)
    add_bullets(doc, [
        "Successful processing rate > 99%.",
        f"SLA compliance achieved (SLA: {v(c.get('integration', {}).get('sla'))}).",
        "No critical production issues post-deployment.",
        "Successful monitoring and alerting operational.",
    ])

    # ── Section 2: Scope ──
    add_heading(doc, "2. Scope", level=1)
    add_heading(doc, "2.1 In Scope", level=2)
    add_bullets(doc, [
        f"Retrieval of {entities_lower} data from {src} via {protocol}.",
        "Data parsing, validation, and transformation.",
        f"Loading of validated data into {tgt}.",
        "Error handling, logging, and notification.",
        "Monitoring and batch status tracking.",
    ])
    add_heading(doc, "2.2 Out of Scope", level=2)
    add_bullets(doc, [v(a) for a in c.get("outOfScope", [{"value": "TBD"}])])

    # ── Section 3: Assumptions ──
    add_heading(doc, "3. Assumptions", level=1)
    a_table = doc.add_table(rows=1, cols=2)
    a_table.style = "Table Grid"
    make_header_row(a_table, ["ID", "Assumption"])
    for i, a in enumerate(c.get("assumptions", [])):
        add_data_row(a_table, [f"A{i+1}", v(a)], shaded=(i % 2 == 1))

    # ── Section 4: Dependencies ──
    add_heading(doc, "4. Dependencies", level=1)
    d_table = doc.add_table(rows=1, cols=3)
    d_table.style = "Table Grid"
    make_header_row(d_table, ["ID", "Dependency", "Owner"])
    for i, d in enumerate(c.get("dependencies", [])):
        add_data_row(d_table, [f"D{i+1}", v(d), "TBD"], shaded=(i % 2 == 1))

    # ── Section 5: Risks ──
    add_heading(doc, "5. Risks and Mitigation", level=1)
    r_table = doc.add_table(rows=1, cols=3)
    r_table.style = "Table Grid"
    make_header_row(r_table, ["Risk", "Impact", "Mitigation"])
    risks = [
        ("Source system unavailable or connectivity failure", "High", "Retry mechanism with alerting; confirm connectivity during setup."),
        ("Invalid or malformed data content", "Medium", "Validation rules applied before loading; rejected records logged."),
        ("Target system connectivity failure", "High", "Retry mechanism; alerting and escalation."),
        ("Authentication credential expiry", "Medium", "Credential rotation procedures; monitoring for auth failures."),
        ("Data format changes without notification", "Medium", "Schema validation at ingestion; monitoring on parse failures."),
    ]
    for i, (risk, impact, mitigation) in enumerate(risks):
        add_data_row(r_table, [risk, impact, mitigation], shaded=(i % 2 == 1))

    # ── Section 6: Solution Overview ──
    add_heading(doc, "6. Solution Overview", level=1)
    add_heading(doc, "6.1 Functional Overview", level=2)
    add_para(doc, f"The integration connects to {src}, retrieves {entities_lower} data, validates and transforms the data per the mapping specification, and loads validated records into {tgt}. Batch status (Completed or Error) is recorded for every run.")
    add_heading(doc, "6.2 Integration Pattern", level=2)
    add_para(doc, pattern)
    add_heading(doc, "6.3 High-Level Architecture", level=2)
    add_para(doc, f"{src} → {middleware} Integration Layer → {tgt}")

    # ── Section 7: Technical Summary ──
    add_heading(doc, "7. Technical Summary", level=1)
    two_col_table(doc, [
        ["Source System", src],
        ["Target System", tgt],
        ["Middleware", middleware],
        ["Trigger Type", v(c.get("integration", {}).get("trigger"))],
        ["Schedule", v(c.get("integration", {}).get("schedule"))],
        ["Protocol", protocol],
        ["Authentication", v(c.get("security", {}).get("sourceAuth"))],
        ["Data Volume", v(c.get("source", {}).get("volume"))],
        ["SLA", v(c.get("integration", {}).get("sla"))],
    ])

    # ── Section 8: Source System ──
    add_heading(doc, "8. Source System Details", level=1)
    two_col_table(doc, [
        ["System Name", src],
        ["Interface Type", f"{v(c['source'].get('interfaceType'))} ({protocol})"],
        ["Authentication Method", v(c.get("security", {}).get("sourceAuth"))],
        ["Data Format", v(c["source"].get("dataFormat"))],
        ["Volume", v(c["source"].get("volume"))],
        ["Frequency", v(c["source"].get("frequency"))],
    ])

    # ── Section 9: Target System ──
    add_heading(doc, "9. Target System Details", level=1)
    two_col_table(doc, [
        ["System Name", tgt],
        ["Interface Type", v(c["target"].get("interfaceType"))],
        ["Authentication Method", v(c.get("security", {}).get("targetAuth"))],
        ["Data Format", v(c["target"].get("dataFormat"))],
        ["Target Table", v(c["target"].get("targetTable"))],
        ["Volume", v(c["target"].get("volume"))],
    ])

    # ── Section 10: End-to-End Flow ──
    add_heading(doc, "10. End-to-End Process Flow", level=1)
    schedule = v(c.get("integration", {}).get("schedule"))
    trigger_text = f"on a {schedule.lower()} schedule" if schedule != "TBD" else "per the configured trigger (TBD)"

    steps = [
        ("Step 1: Trigger", f"The integration is initiated {trigger_text}."),
        ("Step 2: Authentication", f"The process authenticates to {src} using configured credentials (method: {v(c.get('security', {}).get('sourceAuth'))})."),
        ("Step 3: Data Extraction", f"{entities} data is retrieved from {src} and staged for processing."),
        ("Step 4: Validation", "Data is validated against business and technical rules before transformation. Invalid records are rejected and logged."),
        ("Step 5: Transformation", f"Validated records are transformed per the mapping specification ({v(c.get('mappings', {}).get('status'))})."),
        ("Step 6: Processing", "Transformed records are processed in batch and routed to the target system."),
        ("Step 7: Target Update", f"Records are inserted/updated in {tgt} (target table: {v(c['target'].get('targetTable'))})."),
        ("Step 8: Notification", "Success and failure notifications are sent per the notification strategy (Section 21). Batch status is recorded as Completed (C) or Error (E)."),
    ]
    for heading, text in steps:
        add_heading(doc, heading, level=2)
        add_para(doc, text)

    # ── Section 11: Detailed Process Design ──
    add_heading(doc, "11. Detailed Process Design", level=1)
    add_heading(doc, "11.1 Main Integration Process", level=2)
    add_para(doc, f"Purpose: Connect to {src}, retrieve {entities_lower} data, validate, transform, and load into {tgt}.")

    add_heading(doc, "Processing Logic", level=3)
    add_bullets(doc, [
        f"Connect to {src} using configured credentials.",
        f"Retrieve {entities_lower} data.",
        "Validate each record against validation rules (Section 14).",
        "Transform valid records per mapping specification (Section 12).",
        f"Insert/update records in {tgt}.",
        "Log rejected records with error details.",
    ])

    add_heading(doc, "Success Scenario", level=3)
    add_para(doc, f"All valid records are loaded into {tgt}; batch status set to Completed.")
    add_heading(doc, "Failure Scenario", level=3)
    add_para(doc, "Write failure or transformation error is logged; batch status set to Error.")

    # ── Section 12: Data Mapping ──
    add_heading(doc, "12. Data Mapping Specification", level=1)
    add_para(doc, "Field-level mapping is TBD — no mapping specification was provided. The table below will be populated during detailed design.")
    m_table = doc.add_table(rows=1, cols=6)
    m_table.style = "Table Grid"
    make_header_row(m_table, ["Source Field", "Source Type", "Target Field", "Target Type", "Transformation", "Mandatory"])
    add_data_row(m_table, ["TBD", "TBD", "TBD", "TBD", "TBD", "TBD"])

    # ── Section 13: Business Rules ──
    add_heading(doc, "13. Business Rules", level=1)
    br_table = doc.add_table(rows=1, cols=2)
    br_table.style = "Table Grid"
    make_header_row(br_table, ["Rule ID", "Description"])
    for i, r in enumerate(c.get("businessRules", [{"value": "TBD"}])):
        add_data_row(br_table, [f"BR-{i+1:03d}", v(r)], shaded=(i % 2 == 1))

    # ── Section 14: Validation Rules ──
    add_heading(doc, "14. Validation Rules", level=1)
    add_para(doc, "Specific validation rules are TBD pending mapping specification. Standard types:")
    vr_table = doc.add_table(rows=1, cols=3)
    vr_table.style = "Table Grid"
    make_header_row(vr_table, ["Validation Type", "Rule", "Action"])
    for i, (vtype, rule, action) in enumerate([
        ("Mandatory", "TBD — required fields per mapping", "Reject"),
        ("Format", "TBD — data type and format checks", "Reject"),
        ("Duplicate", "TBD — duplicate record detection", "Skip"),
        ("Business", "TBD — business-specific rules", "Error"),
    ]):
        add_data_row(vr_table, [vtype, rule, action], shaded=(i % 2 == 1))

    # ── Section 15: Connectivity ──
    add_heading(doc, "15. Connectivity Specifications", level=1)
    add_heading(doc, "15.1 Source Connectivity", level=2)
    two_col_table(doc, [
        ["System", src],
        ["Protocol", protocol],
        ["Authentication", v(c.get("security", {}).get("sourceAuth"))],
    ])
    doc.add_paragraph()
    add_heading(doc, "15.2 Target Connectivity", level=2)
    two_col_table(doc, [
        ["System", tgt],
        ["Connection Type", "TBD"],
        ["Authentication", v(c.get("security", {}).get("targetAuth"))],
        ["Target Table/Schema", v(c["target"].get("targetTable"))],
    ])

    # ── Section 16: Components ──
    add_heading(doc, "16. Integration Components", level=1)
    add_heading(doc, "Boomi Components", level=2)
    add_para(doc, "Component names will be assigned during build.")
    bc_table = doc.add_table(rows=1, cols=3)
    bc_table.style = "Table Grid"
    make_header_row(bc_table, ["Component Type", "Name", "Purpose"])
    for i, (ctype, purpose) in enumerate([
        ("Process", "Main integration process"),
        ("Connector", f"Source connector for {src}"),
        ("Connector", f"Target connector for {tgt}"),
        ("Map", "Data field mapping"),
        ("Profile", "Source data profile"),
    ]):
        add_data_row(bc_table, [ctype, "TBD", purpose], shaded=(i % 2 == 1))

    doc.add_paragraph()
    add_heading(doc, "Reusable Components", level=2)
    two_col_table(doc, [
        ["Logging Framework", "Standard Boomi logging [STANDARD_DEFAULT]"],
        ["Error Handler", "Standard Boomi error handling [STANDARD_DEFAULT]"],
        ["Notification Service", "TBD"],
    ])

    # ── Section 17: Configuration ──
    add_heading(doc, "17. Configuration Management", level=1)
    add_para(doc, f"Environment-specific properties ({src} credentials, {tgt} connection strings) will be managed via Boomi environment extensions. Details TBD during build.")

    # ── Section 18: Error Handling ──
    add_heading(doc, "18. Error Handling Strategy", level=1)
    add_heading(doc, "Technical Errors", level=2)
    add_bullets(doc, ["Source connectivity or authentication failure", "Target connectivity failure", "Data read/parse failure", "System unavailability"])
    add_heading(doc, "Functional Errors", level=2)
    add_bullets(doc, ["Business rule violation", "Invalid or missing mandatory data", "Data format mismatch"])
    add_heading(doc, "Error Response Matrix", level=2)
    er_table = doc.add_table(rows=1, cols=4)
    er_table.style = "Table Grid"
    make_header_row(er_table, ["Error Type", "Retry", "Notification", "Action"])
    for i, row in enumerate([
        ("Technical", "Yes", "Yes", "Retry per Section 19"),
        ("Functional", "No", "Yes", "Reject record"),
        ("Validation", "No", "Yes", "Archive for review"),
    ]):
        add_data_row(er_table, row, shaded=(i % 2 == 1))

    # ── Section 19: Retry ──
    add_heading(doc, "19. Retry Strategy", level=1)
    rt_table = doc.add_table(rows=1, cols=3)
    rt_table.style = "Table Grid"
    make_header_row(rt_table, ["Error Condition", "Retry Count", "Interval"])
    for i, row in enumerate([
        ("Source connection timeout", "3", "5 minutes"),
        ("Target connection failure", "3", "5 minutes"),
        ("System unavailability", "5", "Exponential backoff"),
    ]):
        add_data_row(rt_table, row, shaded=(i % 2 == 1))

    # ── Section 20: Logging ──
    add_heading(doc, "20. Logging and Monitoring", level=1)
    add_heading(doc, "Logging Requirements", level=2)
    add_bullets(doc, ["Process Name", "Execution ID", "Start / End Time", "Records processed / loaded / failed", "Error details for failed records"])
    add_heading(doc, "Monitoring Requirements", level=2)
    add_bullets(doc, ["Success Rate", "Failure Rate", "Throughput", "Average Execution Time"])

    # ── Section 21: Notifications ──
    add_heading(doc, "21. Notification Strategy", level=1)
    add_heading(doc, "Success", level=2)
    add_para(doc, "Recipients: TBD\nBatch completed successfully. Records loaded: <count>.")
    add_heading(doc, "Failure", level=2)
    add_para(doc, "Recipients: TBD\nBatch failed. Error: <summary>.")

    # ── Section 22: Security ──
    add_heading(doc, "22. Security Design", level=1)
    add_bullets(doc, [
        f"Source authentication: {v(c.get('security', {}).get('sourceAuth'))}",
        f"Target authentication: {v(c.get('security', {}).get('targetAuth'))}",
        f"{v(c.get('security', {}).get('encryption'))} [STANDARD_DEFAULT]",
        "Secure credential storage via Boomi connection configuration [STANDARD_DEFAULT]",
    ])

    # ── Section 23: Retention ──
    add_heading(doc, "23. Archive and Retention", level=1)
    two_col_table(doc, [
        ["Success Records", "TBD"],
        ["Error Records", "TBD"],
        ["Logs", "TBD"],
    ])

    # ── Section 24: Performance ──
    add_heading(doc, "24. Performance Considerations", level=1)
    add_para(doc, f"Expected volume: {v(c.get('source', {}).get('volume'))}. Batch size and performance tuning TBD during build.")

    # ── Section 25: Testing ──
    add_heading(doc, "25. Testing Strategy", level=1)
    add_heading(doc, "Unit Testing", level=2)
    add_para(doc, f"Validate {src} connectivity, data parsing, transformation logic, and {tgt} write operations in isolation.")
    add_heading(doc, "SIT", level=2)
    add_para(doc, f"End-to-end validation with sample data against non-production {src} and {tgt}.")
    add_heading(doc, "UAT", level=2)
    add_para(doc, f"Business validation that loaded records match expected {entities_lower} data.")

    # ── Section 26: Deployment ──
    add_heading(doc, "26. Deployment Strategy", level=1)
    add_heading(doc, "Pre-Deployment Checklist", level=2)
    add_bullets(doc, ["Code review completed", "Unit testing completed", "Source connectivity validated", "Target connectivity validated", "Configuration validated", "Documentation updated"])
    add_heading(doc, "Deployment Steps", level=2)
    add_bullets(doc, ["Package deployment to target Boomi environment", "Environment configuration", "Smoke testing", "Validation of loaded data"])
    add_heading(doc, "Rollback Plan", level=2)
    add_para(doc, "Revert to prior process version; validate during build.")

    # ── Section 27: Support ──
    add_heading(doc, "27. Support Runbook", level=1)
    add_heading(doc, "L1 Support", level=2)
    add_bullets(doc, ["Monitor batch execution status", "Review alerts", "Escalate unresolved errors"])
    add_heading(doc, "L2 Support", level=2)
    add_bullets(doc, ["Root cause analysis", "Reprocessing failed batches", "Configuration review"])
    add_heading(doc, "L3 Support", level=2)
    add_bullets(doc, ["Code fixes", "Enhancement support", "Vendor coordination"])

    # ── Section 28: Operations ──
    add_heading(doc, "28. Operational Procedures", level=1)
    add_bullets(doc, ["Daily: Process status check, failed execution review", "Weekly: Performance review, error trend analysis", "Monthly: SLA compliance, capacity analysis"])

    # ── Section 29: Escalation ──
    add_heading(doc, "29. Escalation Matrix", level=1)
    esc = doc.add_table(rows=1, cols=3)
    esc.style = "Table Grid"
    make_header_row(esc, ["Severity", "Team", "Response Time"])
    for i, row in enumerate([
        ("Critical", "Integration Team", "1 Hour"),
        ("High", "Integration Team", "4 Hours"),
        ("Medium", "Support Team", "1 Business Day"),
        ("Low", "Support Team", "3 Business Days"),
    ]):
        add_data_row(esc, row, shaded=(i % 2 == 1))

    # ── Section 30: Future ──
    add_heading(doc, "30. Future Enhancements", level=1)
    add_bullets(doc, [
        "Define and implement field-level mapping once specifications are confirmed.",
        "Implement email/Slack notifications for success and failure.",
        "Add data reconciliation reporting.",
    ])

    # ── Section 31: References ──
    add_heading(doc, "31. References", level=1)
    add_bullets(doc, [f"BRD — {project_name}", "Mapping Specification (TBD)", "Source System Documentation (TBD)", "Target System Documentation (TBD)"])

    # ── Section 32: Appendix ──
    add_heading(doc, "32. Appendix", level=1)
    add_heading(doc, "Acronyms", level=2)
    acr = doc.add_table(rows=1, cols=2)
    acr.style = "Table Grid"
    make_header_row(acr, ["Acronym", "Description"])
    for i, (a, d) in enumerate([
        ("TDD", "Technical Design Document"), ("BRD", "Business Requirements Document"),
        ("API", "Application Programming Interface"), ("SFTP", "Secure File Transfer Protocol"),
        ("SLA", "Service Level Agreement"), ("ERP", "Enterprise Resource Planning"),
        ("OPCO", "Operating Company"),
    ]):
        add_data_row(acr, [a, d], shaded=(i % 2 == 1))

    doc.add_paragraph()
    add_heading(doc, "Unresolved Items", level=2)
    add_bullets(doc, c.get("unresolvedItems", ["None"]))

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
