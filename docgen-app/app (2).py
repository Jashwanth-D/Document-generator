"""
app.py
DocAgent — Smart Integration Documentation Assistant
Chat interface that detects intent:
  - Generate docs → produces BRD/TDD
  - Ask about docs → answers the question
  - Off-topic → politely declines

v6: Supports file attachments in the chat input (txt, md, docx, pdf).
Attached files are extracted to text and appended to the prompt, so the
existing parse/validate/generate pipeline is untouched.
"""

import streamlit as st
from llm_parser import parse_requirement
from doc_generator import generate_brd, generate_tdd
from flow_generator import generate_flow_drawio
from datetime import datetime
import json
import os
import io
import re
from groq import Groq

# ── Page Config ──
st.set_page_config(
    page_title="DocAgent",
    page_icon="📄",
    layout="centered",
)

# ── Init session state ──
if "messages" not in st.session_state:
    st.session_state.messages = []
if "brd_bytes" not in st.session_state:
    st.session_state.brd_bytes = None
if "tdd_bytes" not in st.session_state:
    st.session_state.tdd_bytes = None
if "flow_drawio" not in st.session_state:
    st.session_state.flow_drawio = None
if "safe_name" not in st.session_state:
    st.session_state.safe_name = "Integration"
if "show_downloads" not in st.session_state:
    st.session_state.show_downloads = False

# ── Styling ──
st.markdown("""
<style>
    .main-header { text-align: center; color: #1F3864; margin-bottom: 0; }
    .sub-header { text-align: center; color: #666; font-size: 14px; margin-top: -10px; margin-bottom: 20px; }
    .scope-notice { background-color: #f0f4ff; padding: 10px 15px; border-radius: 8px; border-left: 4px solid #1F3864; font-size: 13px; color: #333; margin-bottom: 20px; }
    .attachment-chip { display: inline-block; background: #eef2ff; color: #1F3864; padding: 2px 8px; border-radius: 12px; font-size: 12px; margin-right: 6px; }
</style>
""", unsafe_allow_html=True)

# ── Header ──
st.markdown("<h1 class='main-header'>📄 DocAgent</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-header'>Integration Documentation Assistant</p>", unsafe_allow_html=True)
st.markdown("""<div class='scope-notice'>
    💡 I can help you with: <b>generating BRD/TDD documents</b> from requirements,
    <b>answering questions about integration documentation</b>, and
    <b>guiding you on how to write your own</b>. Type your requirement or
    <b>📎 attach a .docx / .pdf / .txt / .md</b> file — I'll read it for you.
</div>""", unsafe_allow_html=True)

# ── API Key ──
api_key = None
try:
    api_key = st.secrets.get("GROQ_API_KEY", None)
except:
    pass

if not api_key:
    with st.sidebar:
        st.markdown("### ⚙️ Configuration")
        api_key = st.text_input("Groq API Key", type="password", help="Get your free key at console.groq.com")
        st.markdown("---")
        st.markdown("**How to get a free API key:**")
        st.markdown("1. Go to [console.groq.com](https://console.groq.com)")
        st.markdown("2. Sign up (free)")
        st.markdown("3. Create an API key")
        st.markdown("4. Paste it above")

with st.sidebar:
    st.markdown("---")
    st.markdown("### 📌 Quick Examples")
    st.markdown("**Generate docs:**")
    st.markdown("_Create an integration that reads PO files from SFTP and loads into ERP database._")
    st.markdown("")
    st.markdown("**Or attach a file:**")
    st.markdown("_Drop a requirement .docx, PDF, or text file straight into the chat._")
    st.markdown("")
    st.markdown("**Ask a question:**")
    st.markdown("_What sections should a TDD have?_")
    st.markdown("_How do I classify work types?_")
    st.markdown("")
    if st.button("🗑️ Clear Chat", use_container_width=True):
        st.session_state.messages = []
        st.session_state.brd_bytes = None
        st.session_state.tdd_bytes = None
        st.session_state.flow_drawio = None
        st.session_state.show_downloads = False
        st.rerun()


# ── Intent classifier + chat system prompt ──
ROUTER_SYSTEM = """You are the intent classifier for DocAgent. Classify each user message into exactly one of: GENERATE, CHAT, DECLINE.

═══ GENERATE ═══
The user wants BRD / TDD / flow documents actually produced. Pick GENERATE when ANY of the following is true:

(a) The message describes data moving between systems — mentions a source AND a target, OR clearly describes moving/syncing/loading/pulling/pushing/reading/writing data.
   Examples that MUST be GENERATE:
   • "Build an integration that reads PO files from SFTP and loads into ERP"
   • "Sync customer accounts from Salesforce to Oracle nightly"
   • "Pull employee data from Workday into JDE"
   • "Push work-order completions from JDE to ServiceNow"
   • "Move invoices from the billing DB to an SFTP server for auditors"
   • "Create an integration that calls the Salesforce REST API to retrieve new customer accounts and inserts them into the Oracle ERP customer master database"

(b) The message explicitly asks to build / create / generate / make / produce / draft documents, BRD, TDD, docs, flow, or diagram.
   Examples that MUST be GENERATE:
   • "Generate the BRD and TDD"
   • "Build me the docs"
   • "Create the flow diagram"
   • "Make the technical design document"
   • "BUILD BRD and TDD and flow diagram"
   • "build me the tdd brd and flow diagram"

(c) The message includes an attached file whose content describes an integration requirement.

If ANY of (a), (b), or (c) applies → GENERATE. Do not second-guess. Do not route to CHAT just because the message is polite, long, or contains extra context.

═══ CHAT ═══
Pick CHAT ONLY when the message is a pure question, greeting, or conceptual discussion with NO integration requirement described AND NO explicit generation request.
   Examples that are CHAT:
   • "Hi" / "Hello" / "How does this work?"
   • "What can you do?"
   • "What's the difference between T1 and T2?"
   • "How should I structure a TDD?"
   • "What are the 32 sections of a TDD?"
   • "Explain the work-type classification"
   • "Can I attach a Word doc?"

═══ DECLINE ═══
Pick DECLINE ONLY when the message is completely unrelated to integration documentation (poems, weather, unrelated coding help, personal advice, math homework).

═══ TIE-BREAKER ═══
If a message could plausibly be either GENERATE or CHAT → pick GENERATE. It is far better to produce documents the user can refine than to keep chatting when they wanted output.

Respond with ONLY this JSON, nothing else:
{"intent": "GENERATE" | "CHAT" | "DECLINE", "reasoning": "one short sentence"}
"""

CHAT_SYSTEM = """You are DocAgent, a friendly and knowledgeable assistant focused on integration documentation.

Your specialty is:
- Generating BRD (Business Requirements Document) and TDD (Technical Design Document) from integration requirements
- Explaining BRD and TDD structure and all 32 TDD sections
- Work type classification (T1: File→DB, T2: API→DB, T3: DB→DB, T4: DB→File, T5: DB→API)
- Integration patterns (batch, real-time, event-driven)
- Best practices for writing integration requirements and documentation
- Anti-hallucination practices (EXPLICIT, DERIVED, STANDARD_DEFAULT, ASSUMPTION, TBD, HUMAN_REQUIRED tags)
- Boomi middleware documentation standards
- Acceptance criteria, scope, assumptions, dependencies, mappings, validation rules
- Testing strategy for integrations

When users ask what you can do, explain your capabilities warmly. When users ask about your knowledge or limits, be honest and helpful about what you cover.

You're friendly and conversational. You can handle greetings, small talk that leads to documentation topics, and meta-questions about yourself. Just gently steer the conversation toward documentation when it drifts too far.

Keep answers concise, practical, and useful. Use examples from integration documentation context when helpful.
"""

DECLINE_MSG = "I appreciate the question! That one's a bit outside what I'm built for though. My specialty is integration documentation — I can generate BRDs and TDDs, explain work types and template structure, or help you figure out how to document your integrations. Want to try any of that?"


# ────────────────────────────────────────────────
#  File extraction — pulls plain text from
#  attached txt / md / docx / pdf files.
# ────────────────────────────────────────────────
def extract_text_from_file(uploaded_file):
    """Return plain text extracted from an UploadedFile, or an error string."""
    name = uploaded_file.name
    lower = name.lower()
    try:
        data = uploaded_file.read()
    except Exception as e:
        return f"[Could not read {name}: {e}]"

    try:
        if lower.endswith((".txt", ".md")):
            return data.decode("utf-8", errors="replace").strip()

        if lower.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            # Also pull table cell text — requirements often live in tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts).strip()

        if lower.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join(pages).strip()

        return f"[Unsupported file type: {name}]"
    except Exception as e:
        return f"[Extraction failed for {name}: {e}]"


def build_combined_prompt(user_text, uploaded_files):
    """
    Merge typed text and extracted file contents into a single prompt
    that the parser can consume. Also returns a display-friendly version
    for the chat history.
    """
    user_text = (user_text or "").strip()
    file_blocks = []
    display_chips = []

    for f in uploaded_files or []:
        content = extract_text_from_file(f)
        display_chips.append(f.name)
        if content and not content.startswith("["):
            file_blocks.append(f"--- Attached file: {f.name} ---\n{content}")
        else:
            file_blocks.append(f"--- Attached file: {f.name} ---\n{content or '[Empty]'}")

    parts = []
    if user_text:
        parts.append(user_text)
    if file_blocks:
        parts.append("\n\n".join(file_blocks))
    combined = "\n\n".join(parts).strip()

    # Display version — keep the user's typed text, list attachments as chips
    if display_chips:
        chips_html = " ".join(
            f"<span class='attachment-chip'>📎 {name}</span>" for name in display_chips
        )
        if user_text:
            display = f"{user_text}\n\n{chips_html}"
        else:
            display = chips_html
    else:
        display = user_text

    return combined, display


def v(field):
    if field is None:
        return "TBD"
    if isinstance(field, str):
        return field
    if isinstance(field, dict) and "value" in field:
        return str(field["value"])
    return "TBD"


# ────────────────────────────────────────────────
#  Deterministic keyword pre-router — bypasses the
#  LLM for obviously-GENERATE messages so classification
#  is 100% reliable for the common cases.
# ────────────────────────────────────────────────
_GEN_VERB_RE = re.compile(
    r'\b(build|create|generate|make|produce|draft|write|prepare|design)\b',
    re.IGNORECASE,
)
_GIVE_ME_RE = re.compile(r'\b(give|show|send)\s+me\b', re.IGNORECASE)
_GEN_TARGET_RE = re.compile(
    r'\b(brd|tdd|doc|docs|document|documents|flow|diagram|diagrams|files?|deliverables?|package)\b',
    re.IGNORECASE,
)
_INTEGRATION_RE = re.compile(r'\bintegration\b', re.IGNORECASE)
_MOVEMENT_RE = re.compile(
    r'\b(pull|push|sync|load|move|read|send|transfer|extract|ingest|import|export|feed|route|copy|fetch|retrieve|writes?)\b',
    re.IGNORECASE,
)
_DIRECTIONAL_RE = re.compile(r'\b(to|from|into|onto)\b', re.IGNORECASE)
_QUESTION_START_RE = re.compile(
    r'^\s*(what|how|why|when|where|which|who|can|is|are|do|does|should|could|would|will|may)\b',
    re.IGNORECASE,
)


def deterministic_intent(msg):
    """
    Return 'GENERATE' if the message clearly asks for document generation.
    Return None for ambiguous cases (let the LLM classify).

    Any one of these fires GENERATE:
      1. Generation verb (build/create/generate/make/...) + target (brd/tdd/doc/flow/diagram/files/...)
      2. Generation verb + the word "integration"
      3. Data-movement verb (pull/push/sync/load/...) + directional word (to/from/into)
      4. "Give me / show me" + target

    Question-word start with no generation verb → None (probably CHAT).
    """
    if not msg or not msg.strip():
        return None

    starts_with_question = bool(_QUESTION_START_RE.match(msg))
    has_gen_verb = bool(_GEN_VERB_RE.search(msg))
    has_give_me = bool(_GIVE_ME_RE.search(msg))
    has_gen_target = bool(_GEN_TARGET_RE.search(msg))

    # Pure question with no generation verb → let LLM handle it
    if starts_with_question and not has_gen_verb and not has_give_me:
        return None

    # Rule 1: gen verb + gen target ("build the BRD", "create the flow diagram")
    if has_gen_verb and has_gen_target:
        return "GENERATE"

    # Rule 2: gen verb + "integration" ("build an integration that ...")
    if has_gen_verb and _INTEGRATION_RE.search(msg):
        return "GENERATE"

    # Rule 3: movement verb + directional preposition ("pull X from Y", "sync A to B")
    if not starts_with_question and _MOVEMENT_RE.search(msg) and _DIRECTIONAL_RE.search(msg):
        return "GENERATE"

    # Rule 4: "give/show me" + a document target ("give me the BRD")
    if has_give_me and has_gen_target:
        return "GENERATE"

    return None


def classify_intent(user_msg, api_key):
    """Classify user message as GENERATE, CHAT, or DECLINE.
    Deterministic keyword pre-router first; LLM only for ambiguous cases."""
    fast = deterministic_intent(user_msg)
    if fast:
        return fast, "Matched deterministic GENERATE trigger"

    client = Groq(api_key=api_key)
    model = os.environ.get("GROQ_MODEL", "openai/gpt-oss-120b")

    # For very long attachments, only send a preview to the router — it
    # only needs to decide intent, not read the whole thing.
    router_input = user_msg if len(user_msg) < 2000 else user_msg[:2000] + "\n[...truncated for intent classification]"

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": ROUTER_SYSTEM},
            {"role": "user", "content": router_input},
        ],
        temperature=0.0,
        max_tokens=100,
    )

    raw = response.choices[0].message.content.strip()
    try:
        raw = re.sub(r'^```json\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        first = raw.find('{')
        last = raw.rfind('}')
        if first != -1 and last != -1:
            result = json.loads(raw[first:last + 1])
            return result.get("intent", "CHAT"), result.get("reasoning", "")
    except:
        pass
    return "CHAT", "Could not classify, defaulting to documentation chat"


def chat_response(user_msg, chat_history, api_key):
    """Generate a documentation-related chat response."""
    client = Groq(api_key=api_key)
    model = os.environ.get("GROQ_MODEL", "openai/gpt-oss-120b")

    messages = [{"role": "system", "content": CHAT_SYSTEM}]
    for msg in chat_history[-10:]:
        if msg["role"] in ("user", "assistant"):
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_msg})

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.3,
        max_tokens=1024,
    )

    return response.choices[0].message.content


# ── Download fragment (must be defined BEFORE the chat history loop below) ──
@st.fragment
def download_section_inline():
    if not st.session_state.brd_bytes:
        return
    safe_name = st.session_state.safe_name
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            label="📘 Download BRD.docx",
            data=st.session_state.brd_bytes,
            file_name=f"BRD_{safe_name}_v1_0.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"brd_dl_{len(st.session_state.messages)}",
        )
    with col2:
        st.download_button(
            label="📗 Download TDD.docx",
            data=st.session_state.tdd_bytes,
            file_name=f"TDD_{safe_name}_v1_0.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"tdd_dl_{len(st.session_state.messages)}",
        )
    with col3:
        if st.session_state.flow_drawio:
            st.download_button(
                label="📊 Download Flow.drawio",
                data=st.session_state.flow_drawio,
                file_name=f"Flow_{safe_name}_v1_0.drawio",
                mime="application/xml",
                use_container_width=True,
                key=f"flow_dl_{len(st.session_state.messages)}",
                help="Skeleton process flow. Import into Lucidchart (Direct file import) or open in draw.io / diagrams.net.",
            )


# ── Display chat history ──
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        # Use HTML for user messages so attachment chips render
        if msg["role"] == "user" and msg.get("has_attachments"):
            st.markdown(msg["content"], unsafe_allow_html=True)
        else:
            st.markdown(msg["content"])

        if msg.get("has_downloads") and st.session_state.brd_bytes:
            download_section_inline()


# ── Chat input with file attachment support ──
chat_value = st.chat_input(
    "Describe your integration requirement or attach a file...",
    accept_file="multiple",
    file_type=["txt", "md", "docx", "pdf"],
)

if chat_value:
    if not api_key:
        st.error("Please enter your Groq API key in the sidebar.")
        st.stop()

    # st.chat_input with accept_file returns a ChatInputValue with .text and .files
    typed_text = getattr(chat_value, "text", "") or ""
    files = getattr(chat_value, "files", None) or []

    if not typed_text.strip() and not files:
        st.stop()

    with st.spinner("Reading attachments..." if files else "Thinking..."):
        combined_prompt, display_msg = build_combined_prompt(typed_text, files)

    # If attachments produced no text at all AND no typed text, bail early
    if not combined_prompt.strip():
        error_msg = "I couldn't read any text from that. Try a different file or type your requirement."
        st.session_state.messages.append({"role": "user", "content": display_msg, "has_attachments": bool(files)})
        st.session_state.messages.append({"role": "assistant", "content": error_msg})
        st.rerun()

    # Show user message
    st.session_state.messages.append({
        "role": "user",
        "content": display_msg,
        "has_attachments": bool(files),
    })
    with st.chat_message("user"):
        st.markdown(display_msg, unsafe_allow_html=True)

    # Classify intent
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            try:
                intent, reasoning = classify_intent(combined_prompt, api_key)
            except Exception as e:
                st.error(f"Error: {str(e)}")
                st.stop()

        # ── GENERATE: produce BRD + TDD ──
        if intent == "GENERATE":
            with st.status("Generating documents...", expanded=True) as status:
                st.write("🤖 Analyzing your requirement...")
                if files:
                    st.write(f"📎 Using content from {len(files)} attachment(s).")

                try:
                    result = parse_requirement(combined_prompt, api_key)
                except Exception as e:
                    error_msg = f"Sorry, I had trouble parsing that requirement: {str(e)}\n\nCould you rephrase it? Make sure to mention the source system, target system, and what data is being moved."
                    st.markdown(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
                    st.stop()

                canonical = result["canonical"]
                validation = result["validation"]
                cross = result["crossValidation"]

                st.write("📋 Generating BRD and TDD...")

                try:
                    brd_buffer = generate_brd(canonical)
                    tdd_buffer = generate_tdd(canonical)
                    flow_drawio = generate_flow_drawio(canonical)
                except Exception as e:
                    error_msg = f"Document generation failed: {str(e)}"
                    st.markdown(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
                    st.stop()

                project_name = v(canonical.get("project", {}).get("name", "Integration"))
                safe_name = "".join(c if c.isalnum() or c in "_ " else "_" for c in project_name)[:50]

                st.session_state.brd_bytes = brd_buffer.getvalue()
                st.session_state.tdd_bytes = tdd_buffer.getvalue()
                st.session_state.flow_drawio = flow_drawio
                st.session_state.safe_name = safe_name

                status.update(label="✅ Documents ready!", state="complete", expanded=True)

            entities = ", ".join(v(e) for e in canonical.get("entities", []))
            summary = f"""**Documents generated successfully!** Here's what I extracted:

| Field | Value |
|-------|-------|
| **Source** | {v(canonical.get('source', {}).get('system'))} |
| **Target** | {v(canonical.get('target', {}).get('system'))} |
| **Work Type** | {v(canonical.get('integration', {}).get('workType'))} |
| **Direction** | {v(canonical.get('integration', {}).get('direction'))} |
| **Pattern** | {v(canonical.get('integration', {}).get('pattern'))} |
| **Entities** | {entities} |

"""
            if cross["status"] == "PASS":
                summary += "✅ **Cross-validation passed** — BRD and TDD are consistent.\n\n"
            else:
                summary += "⚠️ **Review required:**\n"
                for issue in cross["issues"]:
                    summary += f"- {issue}\n"
                summary += "\n"

            tbds = canonical.get("unresolvedItems", [])
            if tbds:
                summary += "**Key TBDs to resolve:**\n"
                for item in tbds[:6]:
                    summary += f"- {item}\n"
                summary += "\n"

            summary += "Download your documents below 👇 (BRD, TDD, and a skeleton process flow importable into [Lucidchart](https://lucid.app) or [draw.io](https://app.diagrams.net))"

            st.markdown(summary)
            st.session_state.messages.append({
                "role": "assistant",
                "content": summary,
                "has_downloads": True,
            })

            download_section_inline()

        # ── CHAT: answer documentation question ──
        elif intent == "CHAT":
            with st.spinner(""):
                try:
                    response = chat_response(combined_prompt, st.session_state.messages, api_key)
                except Exception as e:
                    response = f"Sorry, I encountered an error: {str(e)}"

            st.markdown(response)
            st.session_state.messages.append({"role": "assistant", "content": response})

        # ── DECLINE: out of scope ──
        elif intent == "DECLINE":
            st.markdown(DECLINE_MSG)
            st.session_state.messages.append({"role": "assistant", "content": DECLINE_MSG})


# ── Footer ──
st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: #999; font-size: 12px;'>"
    "DocAgent v12 — AI-powered integration documentation assistant. "
    "Type a requirement, attach a .docx / .pdf / .txt / .md, or ask a question. "
    "Generates BRD, TDD, and a Lucidchart-compatible skeleton process flow. "
    "TBDs highlighted in yellow across all three."
    "</p>",
    unsafe_allow_html=True,
)
